"""One-off, semantically curated conversion of the three supplied DOCX files."""

from __future__ import annotations

import json
import re
import unicodedata
from collections import Counter
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
RAW_DIR = ROOT / "data" / "raw"
OUTPUT_PATH = ROOT / "data" / "processed" / "rag_documents.jsonl"
REPORT_PATH = ROOT / "data" / "reports" / "manual_conversion_report.md"

FIELDS = [
    "id",
    "category",
    "subcategory",
    "title",
    "amac",
    "kapsam",
    "tanim",
    "genel_bilgiler",
    "kosullar",
    "adimlar",
    "istisnalar",
    "surec",
    "standart_yanit",
    "ilgili_dokumanlar",
]

CATEGORY_BY_FILE_AND_NUMBER = {
    "ebrar.docx": {"5": "HESAP_GUVENLIK", "6": "KAMPANYA_PUAN"},
    "esma nur.docx": {"3": "ODEME", "4": "KARGO_TESLIMAT"},
}

PREFIX = {
    "SIPARIS": "ORDER",
    "IADE": "RETURN",
    "ODEME": "PAYMENT",
    "KARGO_TESLIMAT": "SHIPPING",
    "HESAP_GUVENLIK": "ACCOUNT_SECURITY",
    "KAMPANYA_PUAN": "CAMPAIGN_POINT",
}

# IDs are deliberately curated rather than generated from a generic parser.
ID_BY_TITLE = {
    # Sipariş
    "Sipariş Oluşturma": "ORDER_CREATE_001",
    "Sipariş Durumları": "ORDER_STATUS_001",
    "Sipariş İptali": "ORDER_CANCEL_001",
    "Sipariş Adres Değişikliği": "ORDER_ADDRESS_CHANGE_001",
    "Sipariş İçeriği Değişikliği": "ORDER_CONTENT_CHANGE_001",
    "Sipariş Takibi": "ORDER_TRACKING_001",
    "Eksik Ürün Teslimi": "ORDER_MISSING_ITEM_001",
    "Yanlış Ürün Teslimi": "ORDER_WRONG_ITEM_001",
    "Siparişin Bölünmesi": "ORDER_SPLIT_001",
    "Stok Problemi Nedeniyle Sipariş İptali": "ORDER_STOCK_CANCEL_001",
    "Ön Sipariş Ürünleri": "ORDER_PREORDER_001",
    "Fatura Bilgileri ve Sipariş Faturası": "ORDER_INVOICE_001",
    # İade
    "İade Şartları": "RETURN_CONDITIONS_001",
    "İade Süresi": "RETURN_PERIOD_001",
    "İade Talebi Oluşturma": "RETURN_REQUEST_001",
    "İade Kodu Alma": "RETURN_CODE_001",
    "İade Kargo Süreci": "RETURN_SHIPPING_001",
    "Kullanılmış Ürün İadesi": "RETURN_USED_PRODUCT_001",
    "Kusurlu veya Hasarlı Ürün İadesi": "RETURN_DEFECTIVE_DAMAGED_001",
    "İade Reddi Nedenleri": "RETURN_REJECTION_REASONS_001",
    "İade Ücretinin Hesaba Geçmesi": "RETURN_REFUND_PAYMENT_001",
    "Kampanyalı Ürünlerde İade": "RETURN_CAMPAIGN_PRODUCT_001",
    "Kısmi İade": "RETURN_PARTIAL_001",
    "Değişim Talebi": "RETURN_EXCHANGE_REQUEST_001",
    # Ödeme
    "Kartla Ödeme": "PAYMENT_CARD_001",
    "Ödeme Başarısız Hatası": "PAYMENT_FAILED_001",
    "Karttan Para Çekildi Ama Sipariş Oluşmadı": "PAYMENT_CHARGED_ORDER_NOT_CREATED_001",
    "Çift Çekim Problemi": "PAYMENT_DOUBLE_CHARGE_001",
    "Taksitli Ödeme": "PAYMENT_INSTALLMENT_001",
    "Havale / EFT ile Ödeme": "PAYMENT_BANK_TRANSFER_001",
    "Kapıda Ödeme": "PAYMENT_CASH_ON_DELIVERY_001",
    "Para İadesi Süreci": "PAYMENT_REFUND_001",
    "Kupon / Puan Kullanılan Ödemeler": "PAYMENT_COUPON_POINT_001",
    "Ödeme Güvenliği": "PAYMENT_SECURITY_001",
    "Fatura ve Ödeme Uyuşmazlığı": "PAYMENT_INVOICE_MISMATCH_001",
    # Kargo / teslimat
    "Teslimat Süreci": "SHIPPING_DELIVERY_PROCESS_001",
    "Tahmini Teslimat Tarihi": "SHIPPING_ESTIMATED_DATE_001",
    "Kargo Takip İşlemi": "SHIPPING_TRACKING_001",
    "Kargo Firması Seçimi / Değişikliği": "SHIPPING_CARRIER_CHANGE_001",
    "Kargonun Teslim Edilememesi": "SHIPPING_UNDELIVERED_001",
    "Kargo Teslim Edildi Görünüyor Ama Ulaşmadı": "SHIPPING_MARKED_DELIVERED_NOT_RECEIVED_001",
    "Hasarlı Teslimat": "SHIPPING_DAMAGED_DELIVERY_001",
    "Geç Teslimat": "SHIPPING_LATE_DELIVERY_001",
    "Şubeden Teslim Alma": "SHIPPING_BRANCH_PICKUP_001",
    "Adrese Yeniden Teslimat": "SHIPPING_REDELIVERY_001",
    "Birden Fazla Kargo ile Teslimat": "SHIPPING_MULTIPLE_PACKAGES_001",
    # Hesap ve güvenlik
    "Hesap Oluşturma": "ACCOUNT_SECURITY_CREATE_001",
    "Hesap Doğrulama": "ACCOUNT_SECURITY_VERIFY_001",
    "Giriş Yapamama": "ACCOUNT_SECURITY_LOGIN_ISSUE_001",
    "Şifre Sıfırlama": "ACCOUNT_SECURITY_PASSWORD_RESET_001",
    "E-posta Değiştirme": "ACCOUNT_SECURITY_EMAIL_CHANGE_001",
    "Telefon Numarası Değiştirme": "ACCOUNT_SECURITY_PHONE_CHANGE_001",
    "Hesap Silme": "ACCOUNT_SECURITY_DELETE_001",
    "Şüpheli Giriş Bildirimi": "ACCOUNT_SECURITY_SUSPICIOUS_LOGIN_001",
    "İki Aşamalı Doğrulama": "ACCOUNT_SECURITY_TWO_FACTOR_AUTH_001",
    "Bağlı Cihazlar ve Oturum Yönetimi": "ACCOUNT_SECURITY_SESSION_MANAGEMENT_001",
    "Hesap Askıya Alma": "ACCOUNT_SECURITY_SUSPENSION_001",
    "Kişisel Veri Güvenliği": "ACCOUNT_SECURITY_PERSONAL_DATA_001",
    "Yetkisiz Sipariş / Ödeme Bildirimi": "ACCOUNT_SECURITY_UNAUTHORIZED_TRANSACTION_001",
    # Kampanya ve puan
    "Kampanya Kullanımı": "CAMPAIGN_POINT_CAMPAIGN_USE_001",
    "Kupon Kodu Kullanımı": "CAMPAIGN_POINT_COUPON_USE_001",
    "Kupon Geçersiz Hatası": "CAMPAIGN_POINT_COUPON_INVALID_001",
    "Minimum Sepet Tutarı Şartı": "CAMPAIGN_POINT_MINIMUM_CART_AMOUNT_001",
    "Puan Kazanma": "CAMPAIGN_POINT_EARN_POINTS_001",
    "Puan Kullanma": "CAMPAIGN_POINT_USE_POINTS_001",
    "İade Sonrası Puan Durumu": "CAMPAIGN_POINT_RETURN_POINTS_001",
    "Kampanya Koşulları": "CAMPAIGN_POINT_CONDITIONS_001",
    "Kampanyaların Birleştirilmesi": "CAMPAIGN_POINT_COMBINING_001",
    "Süresi Dolan Kupon / Puan": "CAMPAIGN_POINT_EXPIRED_001",
    "Kategoriye Özel Kampanyalar": "CAMPAIGN_POINT_CATEGORY_CAMPAIGNS_001",
}

CURATED_RELATED = {
    "Hesap Oluşturma": ["Hesap Doğrulama", "Giriş Yapamama"],
    "Hesap Doğrulama": ["Hesap Oluşturma", "Giriş Yapamama"],
    "Giriş Yapamama": ["Hesap Doğrulama", "Şifre Sıfırlama", "Hesap Askıya Alma"],
    "Şifre Sıfırlama": ["Giriş Yapamama", "Şüpheli Giriş Bildirimi"],
    "E-posta Değiştirme": ["Hesap Doğrulama", "Şüpheli Giriş Bildirimi"],
    "Telefon Numarası Değiştirme": ["İki Aşamalı Doğrulama", "Giriş Yapamama"],
    "Hesap Silme": ["Kişisel Veri Güvenliği", "Bağlı Cihazlar ve Oturum Yönetimi"],
    "Şüpheli Giriş Bildirimi": ["Şifre Sıfırlama", "Bağlı Cihazlar ve Oturum Yönetimi", "İki Aşamalı Doğrulama"],
    "İki Aşamalı Doğrulama": ["Telefon Numarası Değiştirme", "Şüpheli Giriş Bildirimi"],
    "Bağlı Cihazlar ve Oturum Yönetimi": ["Şüpheli Giriş Bildirimi", "Şifre Sıfırlama"],
    "Hesap Askıya Alma": ["Giriş Yapamama", "Yetkisiz Sipariş / Ödeme Bildirimi"],
    "Kişisel Veri Güvenliği": ["Ödeme Güvenliği", "Hesap Silme"],
    "Yetkisiz Sipariş / Ödeme Bildirimi": ["Şüpheli Giriş Bildirimi", "Ödeme Güvenliği", "Şifre Sıfırlama"],
    "Kampanya Kullanımı": ["Kampanya Koşulları", "Kampanyaların Birleştirilmesi"],
    "Kupon Kodu Kullanımı": ["Kupon Geçersiz Hatası", "Minimum Sepet Tutarı Şartı"],
    "Kupon Geçersiz Hatası": ["Kupon Kodu Kullanımı", "Kampanya Koşulları"],
    "Minimum Sepet Tutarı Şartı": ["Kupon Kodu Kullanımı", "Kampanya Koşulları"],
    "Puan Kazanma": ["Puan Kullanma", "İade Sonrası Puan Durumu"],
    "Puan Kullanma": ["Puan Kazanma", "Süresi Dolan Kupon / Puan"],
    "İade Sonrası Puan Durumu": ["Kısmi İade", "Kampanyalı Ürünlerde İade", "Puan Kazanma"],
    "Kampanya Koşulları": ["Kampanya Kullanımı", "Kampanyaların Birleştirilmesi"],
    "Kampanyaların Birleştirilmesi": ["Kampanya Koşulları", "Kampanya Kullanımı"],
    "Süresi Dolan Kupon / Puan": ["Kupon Kodu Kullanımı", "Puan Kullanma"],
    "Kategoriye Özel Kampanyalar": ["Kampanya Kullanımı", "Kampanya Koşulları"],
    "Kartla Ödeme": ["Ödeme Güvenliği", "Ödeme Başarısız Hatası"],
    "Ödeme Başarısız Hatası": ["Kartla Ödeme", "Karttan Para Çekildi Ama Sipariş Oluşmadı"],
    "Karttan Para Çekildi Ama Sipariş Oluşmadı": ["Ödeme Başarısız Hatası", "Para İadesi Süreci"],
    "Çift Çekim Problemi": ["Para İadesi Süreci", "Ödeme Güvenliği"],
    "Taksitli Ödeme": ["Kartla Ödeme", "Para İadesi Süreci"],
    "Havale / EFT ile Ödeme": ["Ödeme Başarısız Hatası", "Fatura ve Ödeme Uyuşmazlığı"],
    "Kapıda Ödeme": ["Teslimat Süreci", "Para İadesi Süreci"],
    "Para İadesi Süreci": ["İade Ücretinin Hesaba Geçmesi", "Çift Çekim Problemi"],
    "Kupon / Puan Kullanılan Ödemeler": ["Puan Kullanma", "Kupon Kodu Kullanımı", "Kampanyalı Ürünlerde İade"],
    "Ödeme Güvenliği": ["Kişisel Veri Güvenliği", "Yetkisiz Sipariş / Ödeme Bildirimi"],
    "Fatura ve Ödeme Uyuşmazlığı": ["Fatura Bilgileri ve Sipariş Faturası", "Kartla Ödeme"],
    "Teslimat Süreci": ["Tahmini Teslimat Tarihi", "Kargo Takip İşlemi"],
    "Tahmini Teslimat Tarihi": ["Teslimat Süreci", "Geç Teslimat"],
    "Kargo Takip İşlemi": ["Sipariş Takibi", "Kargo Teslim Edildi Görünüyor Ama Ulaşmadı"],
    "Kargo Firması Seçimi / Değişikliği": ["Teslimat Süreci", "Kargo Takip İşlemi"],
    "Kargonun Teslim Edilememesi": ["Adrese Yeniden Teslimat", "Şubeden Teslim Alma"],
    "Kargo Teslim Edildi Görünüyor Ama Ulaşmadı": ["Kargo Takip İşlemi", "Eksik Ürün Teslimi"],
    "Hasarlı Teslimat": ["Kusurlu veya Hasarlı Ürün İadesi", "İade Talebi Oluşturma"],
    "Geç Teslimat": ["Tahmini Teslimat Tarihi", "Kargo Takip İşlemi"],
    "Şubeden Teslim Alma": ["Kargonun Teslim Edilememesi", "Adrese Yeniden Teslimat"],
    "Adrese Yeniden Teslimat": ["Kargonun Teslim Edilememesi", "Şubeden Teslim Alma"],
    "Birden Fazla Kargo ile Teslimat": ["Siparişin Bölünmesi", "Kargo Takip İşlemi"],
}

RELATED_ALIASES = {
    "Kupon veya Puan Kullanılan Ödemeler": "Kupon / Puan Kullanılan Ödemeler",
    "İade Reddı Nedenleri": "İade Reddi Nedenleri",
}

DUPLICATES = [
    "özgür.docx: İade Kodu Alma iki kez yer alıyor; ilk tam kopya kullanıldı.",
    "özgür.docx: Kısmi İade iki kez yer alıyor; ilk tam kopya kullanıldı.",
    "ebrar.docx: İade Sonrası Puan Durumu iki kez yer alıyor; ilk tam kopya kullanıldı.",
    "esma nur.docx: Kargo Firması Seçimi / Değişikliği iki kez yer alıyor; ilk tam kopya kullanıldı.",
]


def clean(text: str) -> str:
    return " ".join(text.replace("\xa0", " ").split()).strip()


def search_key(text: str) -> str:
    return enumize(text).casefold()


def strip_number(text: str) -> str:
    return re.sub(r"^\d+\.\d+\s*", "", clean(text)).strip()


def style_level(style_name: str) -> int:
    match = re.search(r"(\d+)$", style_name)
    return int(match.group(1)) if match else 99


def enumize(text: str) -> str:
    replacements = str.maketrans("ÇĞİÖŞÜçğıöşü", "CGIOSUcgiosu")
    value = unicodedata.normalize("NFKD", text.translate(replacements))
    value = "".join(ch for ch in value if not unicodedata.combining(ch))
    value = re.sub(r"[^A-Za-z0-9]+", "_", value.upper()).strip("_")
    return value


def sentence(text: str) -> str:
    value = clean(text)
    if len(value) >= 2 and value[0] in "“\"" and value[-1] in "”\"":
        value = value[1:-1].strip()
    if not value:
        return ""
    if value[-1] not in ".!?":
        value += "."
    return value


def meaningful_lines(lines: list[str]) -> list[str]:
    result = []
    for line in lines:
        value = clean(line)
        if len(value) >= 2 and value[0] in "“\"" and value[-1] in "”\"":
            value = value[1:-1].strip()
        if not value:
            continue
        low = value.casefold()
        if low in {
            "bu doküman aşağıdaki durumları kapsar:",
            "bu durumda:",
            "aşağıdaki durumlarda:",
            "aşağıdaki durumlarda destek ekibiyle iletişime geçilmelidir:",
            "aşağıdaki durumlarda kullanıcı destek ekibiyle iletişime geçmelidir:",
            "kullanıcı, aşağıdaki kriterlere göre seçim yapabilir:",
            "sistem aşağıdaki kontrolleri gerçekleştirir:",
        }:
            continue
        if value.endswith(":") and len(value.split()) <= 8:
            continue
        result.append(sentence(value))
    return result


def compact(lines: list[str], limit: int = 8) -> str:
    return " ".join(meaningful_lines(lines)[:limit])


def extract_segments(path: Path) -> list[dict]:
    doc = Document(path)
    starts: list[tuple[int, str, str]] = []
    for index, paragraph in enumerate(doc.paragraphs):
        text = clean(paragraph.text)
        style = paragraph.style.name
        if not text:
            continue
        if path.name == "özgür.docx" and style == "Heading 1":
            starts.append((index, text, "SIPARIS" if index < 556 else "IADE"))
        elif path.name == "ebrar.docx":
            if style == "Heading 1" and re.match(r"^[56]\.\d+\s", text):
                number = text.split(".", 1)[0]
                starts.append((index, strip_number(text), CATEGORY_BY_FILE_AND_NUMBER[path.name][number]))
            elif text == "6.2 Kupon Kodu Kullanımı":
                starts.append((index, strip_number(text), "KAMPANYA_PUAN"))
        elif path.name == "esma nur.docx" and style == "Heading 3" and re.match(r"^[34]\.\d+\s", text):
            number = text.split(".", 1)[0]
            starts.append((index, strip_number(text), CATEGORY_BY_FILE_AND_NUMBER[path.name][number]))

    segments = []
    seen_titles: set[str] = set()
    for position, (start, title, category) in enumerate(starts):
        end = starts[position + 1][0] if position + 1 < len(starts) else len(doc.paragraphs)
        if title in seen_titles:
            continue
        seen_titles.add(title)
        paragraphs = []
        for paragraph in doc.paragraphs[start + 1 : end]:
            text = clean(paragraph.text)
            if path.name == "ebrar.docx" and text == "{":
                break
            if path.name == "özgür.docx" and text.startswith("Fine-Tuning Soru Seti"):
                break
            if text:
                paragraphs.append(
                    {
                        "text": text,
                        "style": paragraph.style.name,
                        "level": style_level(paragraph.style.name),
                    }
                )
        segments.append(
            {
                "source": path.name,
                "title": title,
                "category": category,
                "paragraphs": paragraphs,
            }
        )
    return segments


def build_sections(segment: dict) -> list[dict]:
    sections: list[dict] = []
    stack: list[dict] = []
    current: dict | None = None
    for paragraph in segment["paragraphs"]:
        text = paragraph["text"]
        if segment["source"] == "özgür.docx" and text.lower().startswith(
            ("id:", "category:", "subcategory:")
        ):
            continue
        if paragraph["style"].startswith("Heading"):
            level = paragraph["level"]
            while stack and stack[-1]["level"] >= level:
                stack.pop()
            parent = stack[-1]["heading"] if stack else ""
            current = {"heading": text, "level": level, "parent": parent, "lines": []}
            sections.append(current)
            stack.append(current)
        elif current is not None:
            current["lines"].append(text)
    return sections


def matches(section: dict, patterns: tuple[str, ...]) -> bool:
    haystack = search_key(f"{section['parent']} {section['heading']}")
    return any(search_key(pattern) in haystack for pattern in patterns)


def section_lines(sections: list[dict], patterns: tuple[str, ...], direct_only: bool = False) -> list[str]:
    lines: list[str] = []
    for section in sections:
        target = section["heading"] if direct_only else f"{section['parent']} {section['heading']}"
        target_key = search_key(target)
        if any(search_key(pattern) in target_key for pattern in patterns):
            lines.extend(section["lines"])
    return lines


def infer_scope(title: str, category: str) -> str:
    scopes = {
        "SIPARIS": f"NovaCart üzerinden {title.casefold()} konusu ile ilgili sipariş işlemleri için geçerlidir.",
        "IADE": f"NovaCart üzerinden {title.casefold()} konusu kapsamındaki iade ve değişim işlemleri için geçerlidir.",
        "ODEME": f"NovaCart alışverişlerinde {title.casefold()} konusu ile karşılaşan kullanıcılar ve ilgili ödeme işlemleri için geçerlidir.",
        "KARGO_TESLIMAT": f"NovaCart siparişlerinin {title.casefold()} konusu kapsamındaki kargo ve teslimat işlemleri için geçerlidir.",
        "HESAP_GUVENLIK": f"NovaCart hesabında {title.casefold()} işlemini gerçekleştiren veya bu konuda sorun yaşayan kullanıcılar için geçerlidir.",
        "KAMPANYA_PUAN": f"NovaCart alışverişlerinde {title.casefold()} koşullarına tabi kampanya, kupon veya puan işlemleri için geçerlidir.",
    }
    return scopes[category]


def infer_definition(title: str, purpose: str) -> str:
    first = re.split(r"(?<=[.!?])\s+", purpose)[0] if purpose else ""
    if first:
        return first
    return f"{title}, NovaCart üzerindeki ilgili işlem ve kuralların bütününü ifade eder."


def infer_standard_answer(title: str, steps: list[str]) -> str:
    return (
        f"{title} konusunda dokümanda belirtilen işlem adımlarını izleyebilirsiniz. "
        "İşlem tamamlanmıyorsa koşulları kontrol ederek ilgili işlem ayrıntılarıyla destek ekibine başvurabilirsiniz."
    )


def semantic_record(segment: dict) -> tuple[dict, list[str]]:
    title = segment["title"]
    category = segment["category"]
    sections = build_sections(segment)
    used: set[int] = set()
    derived: list[str] = []

    def direct(patterns: tuple[str, ...], limit: int = 8) -> str:
        lines = []
        for i, sec in enumerate(sections):
            if any(search_key(pattern) in search_key(sec["heading"]) for pattern in patterns):
                lines.extend(sec["lines"])
                used.add(i)
        return compact(lines, limit)

    amac = direct(("Amaç",), 4)
    if not amac:
        amac = f"Bu doküman, {title.casefold()} konusundaki koşulları ve izlenecek süreci açıklamak amacıyla hazırlanmıştır."
        derived.append("amac")

    kapsam = direct(("Kapsam",), 8)
    if not kapsam:
        kapsam = infer_scope(title, category)
        derived.append("kapsam")

    tanim = direct(("Tanım",), 4)
    if not tanim:
        for i, sec in enumerate(sections):
            if "nedir" in sec["heading"].casefold() and sec["lines"]:
                tanim = compact(sec["lines"], 3)
                used.add(i)
                break
    if not tanim:
        tanim = infer_definition(title, amac)
        derived.append("tanim")

    standard = direct(("Standart Yanıt",), 4)
    for i, sec in enumerate(sections):
        if search_key("ilgili dokümanlar") in search_key(sec["heading"]):
            used.add(i)

    condition_patterns = (
        "koşul",
        "şart",
        "yapılabilir",
        "yapılamaz",
        "edilebilir",
        "edilemez",
        "hangi durumlarda",
        "gerekli belge",
        "ne zaman",
        "önemli bilgi",
        "dikkat edilmesi",
        "kullanıcının bilmesi",
        "güvenlik önlemleri",
        "kullanıcının işlem yapabileceği",
        "kapsam dışı",
    )
    exception_patterns = (
        "istisna",
        "sık karşılaşılan",
        "destek gerektiren",
        "hata",
        "olası senaryo",
        "olası neden",
        "problem",
        "uyuşmazlık",
        "yanlış teslim",
        "görünmediği",
        "normal kabul edilen",
        "anormal gecikme",
    )
    process_patterns = (
        "sonrası süreç",
        "süreç nasıl işler",
        "değerlendirme süreci",
        "inceleme süreci",
        "ne olur",
        "güncellenmesi",
        "ücret iadesi süreci",
        "kullanım süreci",
        "teslimat yapısı",
        "bekleyen işlemler",
        "verilerin saklanması",
        "yeniden açılması",
        "iade türlerine göre",
        "teslimat süreci üzerindeki etkisi",
        "kargo sürecinde",
        "yeniden teslimat süreci",
        "kargo firması incelemesi",
    )
    action_patterns = (
        "adım",
        "adımları",
        "yapılması gerekenler",
        "nereden",
        "nasıl başlar",
        "nasıl gönderilir",
        "nasıl kontrol edilir",
        "görüntüleme",
        "görüntülenir",
        "etkinleştirme",
        "devre dışı bırakma",
        "tekrar açılır",
        "şubeden teslim alma süreci",
    )

    kosullar: list[str] = []
    istisnalar: list[str] = []
    adimlar: list[str] = []
    surec_lines: list[str] = []

    for i, sec in enumerate(sections):
        heading = sec["heading"]
        lines = meaningful_lines(sec["lines"])
        if not lines:
            continue
        if matches(sec, exception_patterns):
            istisnalar.extend(lines)
            used.add(i)
        elif matches(sec, condition_patterns):
            kosullar.extend(lines)
            used.add(i)
        elif matches(sec, process_patterns):
            surec_lines.extend(lines)
            used.add(i)
        elif matches(sec, action_patterns):
            if re.match(r"^Adım\s+\d+", heading, re.IGNORECASE):
                label = re.sub(r"^Adım\s+\d+\s*:?\s*", "", heading, flags=re.IGNORECASE)
                combined = " ".join(lines)
                adimlar.append(sentence(f"{label}: {combined}" if label else combined))
            else:
                adimlar.extend(lines)
            used.add(i)

    # Explicit "İstisnalar" always wins even if a parent heading also resembles a process.
    explicit_exceptions = section_lines(sections, ("İstisnalar",), direct_only=True)
    for item in meaningful_lines(explicit_exceptions):
        if item not in istisnalar:
            istisnalar.append(item)

    # If no formal action section exists, use ordered process statements as user-facing steps.
    if not adimlar:
        candidates = section_lines(
            sections,
            ("Süreç", "İşlem", "Talep", "Kullanıcının Yapması Gerekenler"),
            direct_only=True,
        )
        adimlar = meaningful_lines(candidates)[:8]

    if not surec_lines:
        candidates = section_lines(
            sections,
            ("Sonrası", "Süreç", "İnceleme", "Değerlendirme", "Tamamlanması"),
            direct_only=True,
        )
        surec_lines = meaningful_lines(candidates)

    # General information is selected from the remaining descriptive sections.
    general_lines: list[str] = []
    descriptive_patterns = (
        "genel bilgi",
        "nedir",
        "açıklaması",
        "durumları",
        "türleri",
        "avantaj",
        "dezavantaj",
        "güvenlik avantajları",
        "toplanan veriler",
        "nasıl korunur",
        "şifre güvenliği",
        "kart ve ödeme bilgileri",
        "kimlerle paylaşılır",
    )
    for i, sec in enumerate(sections):
        if i in used or not sec["lines"]:
            continue
        if matches(sec, descriptive_patterns):
            general_lines.extend(meaningful_lines(sec["lines"]))
            used.add(i)
    if not general_lines:
        for i, sec in enumerate(sections):
            if i not in used and sec["lines"]:
                general_lines.extend(meaningful_lines(sec["lines"]))
            if len(general_lines) >= 6:
                break
    genel_bilgiler = " ".join(general_lines[:8])
    if not genel_bilgiler:
        genel_bilgiler = tanim
        derived.append("genel_bilgiler")

    if not standard:
        standard = infer_standard_answer(title, adimlar)
        derived.append("standart_yanit")

    explicit_related = []
    for sec in sections:
        if search_key("ilgili dokümanlar") in search_key(sec["heading"]):
            explicit_related.extend(clean(line) for line in sec["lines"] if clean(line))

    related_titles = explicit_related or CURATED_RELATED.get(title, [])
    related_ids = []
    unresolved = []
    for related_title in related_titles:
        canonical = RELATED_ALIASES.get(related_title, related_title)
        if canonical in ID_BY_TITLE:
            related_ids.append(ID_BY_TITLE[canonical])
        else:
            unresolved.append(related_title)
    if not explicit_related:
        derived.append("ilgili_dokumanlar")

    record = {
        "id": ID_BY_TITLE[title],
        "category": category,
        "subcategory": enumize(title),
        "title": title,
        "amac": amac,
        "kapsam": kapsam,
        "tanim": tanim,
        "genel_bilgiler": genel_bilgiler,
        "kosullar": list(dict.fromkeys(kosullar)),
        "adimlar": list(dict.fromkeys(adimlar)),
        "istisnalar": list(dict.fromkeys(istisnalar)),
        "surec": " ".join(dict.fromkeys(surec_lines)),
        "standart_yanit": standard,
        "ilgili_dokumanlar": list(dict.fromkeys(related_ids)),
    }
    return record, derived + [f"unresolved:{item}" for item in unresolved]


def validate(records: list[dict]) -> None:
    allowed = {"SIPARIS", "IADE", "ODEME", "KARGO_TESLIMAT", "HESAP_GUVENLIK", "KAMPANYA_PUAN"}
    ids = []
    subcategories = []
    for index, record in enumerate(records, start=1):
        assert list(record) == FIELDS, f"Satır {index}: alan sırası/şeması hatalı"
        assert record["category"] in allowed, f"Satır {index}: geçersiz kategori"
        for field in ("kosullar", "adimlar", "istisnalar", "ilgili_dokumanlar"):
            assert isinstance(record[field], list), f"Satır {index}: {field} liste değil"
            assert all(isinstance(item, str) for item in record[field]), f"Satır {index}: {field} elemanı string değil"
        for field in set(FIELDS) - {"kosullar", "adimlar", "istisnalar", "ilgili_dokumanlar"}:
            assert isinstance(record[field], str), f"Satır {index}: {field} string değil"
        ids.append(record["id"])
        subcategories.append(record["subcategory"])
    assert len(ids) == len(set(ids)), "Tekrarlanan ID var"
    assert len(subcategories) == len(set(subcategories)), "Tekrarlanan subcategory var"
    assert len(records) == 70, f"Beklenen 70 kayıt yerine {len(records)} üretildi"


def write_report(records: list[dict], source_count: int, notes_by_title: dict[str, list[str]]) -> None:
    derived_counts = Counter(
        note
        for notes in notes_by_title.values()
        for note in notes
        if not note.startswith("unresolved:")
    )
    unresolved = [
        f"- **{title}:** {note.split(':', 1)[1]}"
        for title, notes in notes_by_title.items()
        for note in notes
        if note.startswith("unresolved:")
    ]

    manual_review = [
        "- Kaynakta yinelenen kayıtlar tekilleştirildi: İade Kodu Alma, Kısmi İade, İade Sonrası Puan Durumu ve Kargo Firması Seçimi / Değişikliği.",
        "- `6.2 Kupon Kodu Kullanımı` kaynakta başlık stili olmadan yazılmıştı; içerik ayrı doküman olarak değerlendirildi.",
        "- Hesap/güvenlik ile ödeme/kargo belgelerinde `Kapsam`, `Tanım`, `Standart Yanıt` ve `İlgili Dokümanlar` alanları çoğunlukla açık başlıklarla verilmediği için içerikten kısa ve genel ifadeler türetildi.",
        "- Kaynaktaki süre, limit, güvenlik ve politika ifadeleri aynen korunmalı; üretim ortamına alınmadan önce iş birimi tarafından doğrulanması önerilir.",
    ]

    missing_lines = [
        f"- `{field}`: {count} kayıtta kaynakta doğrudan alan bulunmadığı için bağlamdan türetildi."
        for field, count in sorted(derived_counts.items())
    ]
    if not missing_lines:
        missing_lines = ["- Kaynakta eksik olup bağlamdan türetilen alan bulunmadı."]

    report = f"""# Manuel DOCX Dönüşüm Raporu

## Özet

- Okunan DOCX sayısı: **{source_count}**
- Üretilen benzersiz doküman sayısı: **{len(records)}**
- Çıktı: `data/processed/rag_documents.jsonl`

## Eksik veya açıkça etiketlenmemiş alanlar

{chr(10).join(missing_lines)}

Boş bırakılan zorunlu alanlar, kaynakta ilgili bilgi bulunmayan liste alanlarıdır. String alanlar, yalnızca kaynak bağlamının güvenli biçimde desteklediği ölçüde tamamlandı.

## ID ile eşleştirilemeyen ilgili dokümanlar

{chr(10).join(unresolved) if unresolved else "- Tüm belirtilen ilgili dokümanlar üretilen kayıt ID'leriyle eşleştirildi."}

## Kaynak tekrarları ve biçim sorunları

{chr(10).join(f"- {item}" for item in DUPLICATES)}

## Manuel kontrol gerektiren kayıtlar

{chr(10).join(manual_review)}

## Doğrulama sonucu

- Her JSONL satırı geçerli JSON olarak doğrulandı.
- Tüm zorunlu alanlar mevcut.
- Kategori değerleri izin verilen enumlarla sınırlı.
- ID ve subcategory değerleri benzersiz.
- Liste ve string alan tipleri şemaya uygun.
"""
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    REPORT_PATH.write_text(report, encoding="utf-8")


def main() -> None:
    paths = sorted(RAW_DIR.glob("*.docx"))
    segments = []
    for path in paths:
        segments.extend(extract_segments(path))

    missing_id_titles = sorted({segment["title"] for segment in segments} - set(ID_BY_TITLE))
    if missing_id_titles:
        raise ValueError(f"ID eşlemesi eksik başlıklar: {missing_id_titles}")

    records = []
    notes_by_title = {}
    for segment in segments:
        record, notes = semantic_record(segment)
        records.append(record)
        notes_by_title[record["title"]] = notes

    category_order = {
        "SIPARIS": 0,
        "IADE": 1,
        "ODEME": 2,
        "KARGO_TESLIMAT": 3,
        "HESAP_GUVENLIK": 4,
        "KAMPANYA_PUAN": 5,
    }
    records.sort(key=lambda item: (category_order[item["category"]], item["id"]))
    validate(records)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    with OUTPUT_PATH.open("w", encoding="utf-8", newline="\n") as handle:
        for record in records:
            handle.write(json.dumps(record, ensure_ascii=False, separators=(",", ":")) + "\n")

    # Re-read the exact artifact to ensure every physical line is valid JSON.
    parsed = [json.loads(line) for line in OUTPUT_PATH.read_text(encoding="utf-8").splitlines()]
    validate(parsed)
    write_report(parsed, len(paths), notes_by_title)
    print(f"DOCX={len(paths)} RECORDS={len(parsed)} OUTPUT={OUTPUT_PATH}")


if __name__ == "__main__":
    main()
